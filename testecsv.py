from docx import Document           #import para os arquivos docx
from docx.shared import Pt          #import para os funcoes docx
from datetime import date           #import para as datas
from tkinter import messagebox      #import para o alerta
from tkinter import simpledialog    #digitar informações
from tkinter.messagebox import askyesnocancel
import socket                       #import para informacoes da maquina
import locale                       #import para localidade
import csv

#Para identificar a linguagem/local(ptBR)
locale.setlocale(locale.LC_ALL, 'pt_BR.utf-8')

#Para identificar que maquina realizou o requerimento
ip_local = socket.gethostbyname(socket.gethostname())
host_name = socket.gethostname()

data_atual = date.today()

encontrouMatricula = None

def Abonada(abono):###############################################
    # Use barras invertidas duplicadas ou uma barra invertida crua para evitar caracteres de escape
    # Caminho do arquivo
    document = Document(r'C:\Users\Usuário\Desktop\Bruno\abono\ABONADAteste.docx')


    #escrevendo no final do documento 
    #p = document.add_paragraph(f'Documento expedido pela maquina: NOME DO HOST:{host_name} IP LOCAL:{ip_local}')

    #locais do texto no arquivo que irá modificar
    m = 'MatriculaDoServidor'
    s = 'NomeDoServidor'
    c = 'CargoDoServidor'
    d = 'DataDoAbono'
    hoje = "DataDoDocumento"

    # Iterar sobre todos os parágrafos do documento
    for paragraph in document.paragraphs:
        # Configurar a fonte para cada execução no parágrafo
        for run in paragraph.runs:
            font = run.font
            font.name = 'Arial'
            font.size = Pt(10)

        # Substituir as variáveis nos parágrafos
        if s in paragraph.text:
            paragraph.text = paragraph.text.replace(s, nomeS)
        if c in paragraph.text:
            paragraph.text = paragraph.text.replace(c, cargo)
        if m in paragraph.text:
            paragraph.text = paragraph.text.replace(m, matr)
        if d in paragraph.text:
            dataAbono = abono.strftime('%d/%m/%Y')
            paragraph.text = paragraph.text.replace(d, dataAbono)
        if hoje in paragraph.text:
            data_em_texto = data_atual.strftime('%d de %B de %Y')
            paragraph.text = paragraph.text.replace(hoje, data_em_texto)

        print(paragraph.text)

    #Salvar documento
    document.save(r'C:\Users\Usuário\Desktop\Bruno\abono\ABONADAteste1.docx')
##################################################################
def buscaMatr(arq):###########EM MANUTENÇÃO############
    answer1 = simpledialog.askstring("Input", "Qual sua matricula?")
    if answer1 is not None:
        for linha in enumerate(arquivo_csv):
            if matr == linha[6]:
                dataAbonada()
                encontrouMatricula = True
                break
            else:
                print("Matricula não encontrada")
                encontrouMatricula = False
    elif answer1 is None:
        buscaMatr(answer1)
    for linha in arq:
        if matricula == linha[6]:
            return True
        else: return False
##################################################################
def dataAbonada():################################################
    diaAbono = simpledialog.askinteger("Input", "Qual dia pretende abonar?")
    if diaAbono is not None:
        mesAbono = simpledialog.askinteger("Input", "De qual mês?")
        if mesAbono is not None:
            anoAbono = simpledialog.askinteger("Input", "Ano?")
            if anoAbono is not None:
                #convertendo para data
                abono = date(anoAbono,mesAbono,diaAbono)
                #verificando proximidade das datas
                if abono.toordinal() < (data_atual.toordinal()+7):
                    alertas(abono)
                else:
                    Abonada(abono)
##################################################################
def alertas(data):################################################
    #alerta
    answer = askyesnocancel(
        title = 'Pedido muito proximo do abono', 
        message = 'A abonada deve ser requerida com no minimo 7(sete) dias de antecedencia. \nPressione SIM para proseguir com o Abono \nPressione NÃO para colocar outra data \nPressione CANCELAR para finalizar o programa')
    if answer == True:
        Abonada(data)
    elif answer == False:
        dataAbonada()
    else:
        messagebox.showinfo('Pedido muito proximo do abono', \
            'Finalizando programa')
        print("Fim")
##################################################################

#abre base da dados csv
with open("C://Users/Usuário/Desktop/Bruno/abono/planilha SERVIÇOS URBANOS.csv","r", encoding='utf-8') as arquivo:
    arquivo_csv = csv.reader(arquivo, delimiter=',')
    #Usuario entrega matricula
    answer1 = simpledialog.askstring("Input", "Qual sua matricula?")
    if answer1 is not None:
        matr = answer1
        for i, linha in enumerate(arquivo_csv):
            if matr == linha[6]:
                nomeS = linha[1]
                cargo = linha[7]
                print("Nome: " + nomeS)
                print("Cargo: " + cargo)
                dataAbonada()
                encontrouMatricula = True
                break
            else:
                print("Matricula não encontrada")
                encontrouMatricula = False
    

    
    
    if encontrouMatricula == False:
        print("Digite uma matricula existente")

            